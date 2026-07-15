"""Extract raw text from uploaded resumes and structure it with the LLM."""

import io

from .. import llm, prompts


class ResumeParseError(Exception):
    pass


def extract_text(filename: str, content: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _extract_pdf(content)
    if name.endswith(".docx"):
        return _extract_docx(content)
    if name.endswith((".txt", ".md", ".text")):
        return content.decode("utf-8", errors="replace")
    # Fall back on content sniffing
    if content[:5] == b"%PDF-":
        return _extract_pdf(content)
    if content[:2] == b"PK":  # docx is a zip
        return _extract_docx(content)
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError as e:
        raise ResumeParseError(
            "Unsupported resume format. Upload a PDF, DOCX, or plain-text file, "
            "or paste the resume text directly."
        ) from e


def _extract_pdf(content: bytes) -> str:
    import pdfplumber

    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
    except Exception as e:
        raise ResumeParseError(f"Could not read PDF: {e}") from e
    text = "\n\n".join(pages).strip()
    if not text:
        raise ResumeParseError(
            "No selectable text found in this PDF (it may be a scanned image). "
            "Please paste the resume text instead."
        )
    return text


def _extract_docx(content: bytes) -> str:
    import docx

    try:
        document = docx.Document(io.BytesIO(content))
    except Exception as e:
        raise ResumeParseError(f"Could not read DOCX: {e}") from e
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    text = "\n".join(part for part in parts if part.strip()).strip()
    if not text:
        raise ResumeParseError("The DOCX file appears to be empty.")
    return text


def parse_resume(resume_text: str) -> dict:
    """Structure the resume text into the profile schema via the LLM."""
    if len(resume_text.strip()) < 50:
        raise ResumeParseError("The resume text is too short to parse.")
    return llm.structured(
        system=prompts.RESUME_SYSTEM,
        user=f"Resume text:\n\n{resume_text[:60000]}",
        schema=prompts.RESUME_SCHEMA,
        max_tokens=8192,
    )
